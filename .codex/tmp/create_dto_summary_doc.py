from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "/Users/m4/Desktop/天之灵项目说明书里面/天之灵项目DTO小白说明书.docx"


def set_east_asia_font(run, font_name="STHeiti"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_east_asia_font(r)
    r.font.size = Pt(10.5)
    r.bold = bold


def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_east_asia_font(r)
    r.font.size = Pt(11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    set_east_asia_font(r)
    r.font.color.rgb = RGBColor(0, 0, 0)
    r.bold = False
    if level == 1:
        r.font.size = Pt(20)
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        r.font.size = Pt(16)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    else:
        r.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_east_asia_font(r)
    r.font.size = Pt(11)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    rt = p.add_run(title)
    set_east_asia_font(rt)
    rt.bold = True
    rt.font.size = Pt(11)
    rb = p.add_run("\n" + body)
    set_east_asia_font(rb)
    rb.font.size = Pt(10.5)
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "STHeiti"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "STHeiti")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "STHeiti")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("天之灵项目 DTO 小白说明书")
    set_east_asia_font(r)
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    rs = subtitle.add_run("从整体业务角色理解 DTO，而不是死记技术参数")
    set_east_asia_font(rs)
    rs.font.size = Pt(11)
    rs.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "一、DTO 在项目里是什么", 1)
    add_paragraph(
        doc,
        "在天之灵项目里，DTO 可以理解成前端和后端沟通时的数据约定。当前端要登录、创建智能体、发送消息、下单、上传文件时，都会把一份数据交给后端。DTO 的作用，就是说明这份数据属于哪个业务、代表什么意图、后端应该怎样接收。",
    )
    add_paragraph(
        doc,
        "DTO 不是业务本身。它更像业务入口处的说明牌：它帮助 controller 看懂请求，也帮助 service 拿到相对稳定、清楚的数据。",
    )
    add_callout(
        doc,
        "小白记忆法",
        "Controller 负责接电话，DTO 负责确认对方说的是哪件事，Service 负责真正去办事。",
    )

    add_heading(doc, "二、DTO 在整体流程中的位置", 1)
    add_paragraph(doc, "一个典型请求在项目里的流动方式是：")
    for item in [
        "前端页面发起请求。",
        "Controller 接收请求，并把请求体交给对应 DTO。",
        "DTO 描述或校验这次请求携带的数据。",
        "Service 根据这些数据执行业务逻辑。",
        "后端把处理结果返回给前端。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "三、各类 DTO 的业务分工", 1)
    rows = [
        ("user.dto.ts", "用户是谁", "登录、验证码、微信小程序登录、绑定手机号、修改个人资料。"),
        ("agent.dto.ts", "智能体是谁", "创建和维护数字亲人/联系人档案，包括头像、默认联系人和人物资料。"),
        ("conversation.dto.ts", "怎么互动", "发送聊天消息、语音转文字、生成纪念照片，支撑用户和智能体之间的互动。"),
        ("post.dto.ts", "怎么发动态", "发布帖子、发表评论、回复评论，并支持提醒智能体参与社区互动。"),
        ("order.dto.ts", "怎么买服务", "购买 VIP 套餐和声音包，把用户购买意图转换成订单请求。"),
        ("storage.dto.ts", "文件怎么上传", "申请云存储上传地址，支撑头像、图片、语音等资源上传。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(1.45), Inches(1.35), Inches(3.7)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    header = table.rows[0].cells
    for idx, text in enumerate(["DTO 文件", "业务问题", "重点功能角色"]):
        set_cell_text(header[idx], text, bold=True)
        set_cell_shading(header[idx], "F2F4F7")
    for file_name, question, role in rows:
        cells = table.add_row().cells
        for idx, text in enumerate([file_name, question, role]):
            set_cell_text(cells[idx], text)
    doc.add_paragraph()

    sections = [
        (
            "四、用户 DTO：管理账号入口",
            "user.dto.ts 负责用户账号体系相关的入口数据。它让后端知道用户是在申请验证码、手机号登录、微信登录，还是在修改昵称、头像、性别、地区和偏好。",
            "它的业务重点不是保存用户资料，而是把用户身份相关操作分清楚，让登录和资料修改流程有明确入口。",
        ),
        (
            "五、智能体 DTO：管理数字亲人档案",
            "agent.dto.ts 负责智能体资料管理。创建智能体时，它承接名字、性别、双方称呼这些基础关系信息；更新资料时，它承接生日、生平、性格、语言习惯、爱好、共同回忆等人物设定。",
            "它的业务重点是支撑“这个数字亲人是谁”。聊天、纪念照片、声音服务等后续功能，都要建立在这份智能体档案之上。",
        ),
        (
            "六、会话 DTO：支撑用户和智能体互动",
            "conversation.dto.ts 负责聊天会话相关入口。用户发送文字、语音、图片，或者请求语音转文字、生成纪念照片时，都会经过这一类 DTO。",
            "它的业务重点是支撑“用户怎样和智能体互动”。如果 agent DTO 管人物档案，conversation DTO 管的就是围绕这个人物发生的对话和回忆生成。",
        ),
        (
            "七、帖子 DTO：支撑动态和评论",
            "post.dto.ts 负责社区动态相关入口。它支撑发布帖子、带图片发动态、评论帖子、回复评论，以及提醒智能体参与回复。",
            "它的业务重点是把一对一聊天之外的互动扩展到社区场景，让用户可以通过动态和评论表达内容，也让智能体有机会参与到公开互动里。",
        ),
        (
            "八、订单 DTO：承接购买意图",
            "order.dto.ts 负责购买和支付入口。用户购买 VIP 套餐或声音包时，订单 DTO 会告诉后端用户想买什么、为哪个智能体购买，以及这次支付需要哪些小程序相关信息。",
            "它的业务重点是把“我想买这个服务”转换成后端可以创建、支付、追踪的订单。",
        ),
        (
            "九、存储 DTO：支撑图片和语音上传",
            "storage.dto.ts 负责上传前的签名申请。头像、帖子图片、聊天语音、纪念照片相关素材，都可能需要先通过它拿到云存储上传地址。",
            "它的业务重点是作为媒体资源的入口基础设施。很多业务看起来在处理头像、图片、语音，底层都离不开文件上传能力。",
        ),
    ]
    for heading, body, point in sections:
        add_heading(doc, heading, 1)
        add_paragraph(doc, body)
        add_callout(doc, "重点角色", point)

    add_heading(doc, "十、学习 DTO 时最该抓住什么", 1)
    for item in [
        "先看这个 DTO 属于哪个业务模块，不要一上来死抠字段细节。",
        "再看它服务哪个接口，理解前端在什么场景下会传这份数据。",
        "最后看 service 怎么使用它，理解这份数据最后推动了什么业务动作。",
    ]:
        add_bullet(doc, item)
    add_paragraph(
        doc,
        "对小白开发者来说，DTO 最重要的价值是帮你建立项目地图：用户、智能体、会话、帖子、订单、存储这些模块通过 DTO 把业务入口分开，项目因此更容易理解和维护。",
    )

    doc.save(OUT)


if __name__ == "__main__":
    main()
